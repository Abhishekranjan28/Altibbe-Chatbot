import streamlit as st
import os
import tempfile
import io
from google.generativeai import configure, GenerativeModel
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain.schema import Document as LangChainDocument
from fpdf import FPDF
import urllib.parse
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

# Load environment variables
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if api_key:
    configure(api_key=api_key)
    model = GenerativeModel("gemini-1.5-flash")
else:
    st.error("Gemini API key not found. Please set it in the .env file.")

# Function to load documents
def load_documents(file):
    try:
        if file.type == "application/pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                tmp_file.write(file.getvalue())
                tmp_file_path = tmp_file.name
            loader = PyPDFLoader(tmp_file_path)
            return loader.load()
        elif file.type == "text/plain":
            text = file.getvalue().decode("utf-8")
            return [LangChainDocument(page_content=text)]
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                tmp_file.write(file.getvalue())
                tmp_file_path = tmp_file.name
            docx = DocxDocument(tmp_file_path)
            text = "\n".join([para.text for para in docx.paragraphs])
            return [LangChainDocument(page_content=text)]
        else:
            st.error(f"Unsupported file format: {file.type}")
            return []
    except Exception as e:
        st.error(f"Error processing file {file.name}: {e}")
        return []

# Create FAISS index
def create_faiss_index(documents):
    embeddings = HuggingFaceEmbeddings()
    return FAISS.from_documents(documents, embeddings)

# Query Gemini model
def query_gemini(prompt):
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {e}"

# Save chat history as PDF
def save_chat_as_pdf(chat_history):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Add the Roboto font (ensure the font file paths are correct)
    pdf.add_font('Roboto', '', './Roboto-Regular.ttf', uni=True)
    pdf.add_font('Roboto-Bold', 'B', './Roboto-Bold.ttf', uni=True)
    pdf.set_font('Roboto', '', 12)

    # Title of the document
    pdf.cell(200, 10, txt="Chat History", ln=True, align="C")
    pdf.ln(10)

    # Add chat history to the PDF
    for chat in chat_history:
        role = "You" if chat['role'] == 'user' else "Bot"
        pdf.multi_cell(0, 10, txt=f"{role}: {chat['message']}")
        pdf.ln(2)

    # Output the generated PDF file to a BytesIO buffer
    pdf_output = io.BytesIO()
    pdf_data = pdf.output(dest='S').encode('latin1')  # Output as a string and encode
    pdf_output.write(pdf_data)
    pdf_output.seek(0)
    return pdf_output

# Create DOCX file
def create_docx(content):
    doc = DocxDocument()
    for chat in content:
        line = f"{'You' if chat['role'] == 'user' else 'Bot'}: {chat['message']}"
        doc.add_paragraph(line)
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# Generate WhatsApp shareable link
def generate_whatsapp_link(content):
    base_url = "https://api.whatsapp.com/send?text="
    message = "\n".join(
        [f"{'You' if chat['role'] == 'user' else 'Bot'}: {chat['message']}" for chat in content]
    )
    encoded_message = urllib.parse.quote(message)
    return f"{base_url}{encoded_message}"

# Send email via SMTP
def send_email_via_smtp(content, recipient_email):
    sender_email = "abhisshekranjan28@gmail.com"
    sender_password = "zkka ewdm tvat zrdh"
    smtp_server = "smtp.gmail.com"
    smtp_port = 587
    try:
        message = MIMEMultipart()
        message["From"] = sender_email
        message["To"] = recipient_email
        message["Subject"] = "Chat History from ChatBot"
        message.attach(MIMEText(content, "plain"))
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(message)
        st.success(f"Email sent to {recipient_email} successfully!")
    except Exception as e:
        st.error(f"Error sending email: {e}")

# Streamlit app
st.set_page_config(page_title="Chatbot-Report Generator", layout="wide")

col_left, col_right = st.columns([3, 3])

with col_left:
    st.image("Altibbe logo dark.png", width=150)

with col_right:
    st.image("Hedamo.jpg", width=200)
    
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

st.title("Report Generator and Chatbot")
st.write("Upload documents and chat with the bot.")

# File uploader
uploaded_files = st.file_uploader("Upload Files", type=["pdf", "txt", "docx"], accept_multiple_files=True)

# Process uploaded files
if uploaded_files:
    documents = []
    for file in uploaded_files:
        file_documents = load_documents(file)
        if file_documents:
            documents.extend(file_documents)
    if documents:
        index = create_faiss_index(documents)
        st.success("Documents successfully loaded and indexed.")

# Chat interface
st.subheader("Chat Interface")
for chat in st.session_state.chat_history:
    role = "🧑‍💻 You" if chat['role'] == 'user' else "🤖 Bot"
    st.markdown(f"**{role}:** {chat['message']}")

user_input = st.text_area("Type your message:", key="user_input", height=100)

if st.button("Send"):
    if user_input.strip() != "":
        st.session_state.chat_history.append({"role": "user", "message": user_input})

        # Generate bot response
        if uploaded_files:
            retriever = index.as_retriever()
            relevant_docs = retriever.get_relevant_documents(user_input)
            context = "\n\n".join(set(doc.page_content for doc in relevant_docs))
            prompt = f"Context: {context}\n\nUser: {user_input}\nBot:"
            bot_response = query_gemini(prompt)
        else:
            bot_response = query_gemini(user_input)

        st.session_state.chat_history.append({"role": "bot", "message": bot_response})
        st.rerun()

# Download buttons
if st.session_state.chat_history:
    st.download_button(
        "Download Chat History as PDF",
        save_chat_as_pdf(st.session_state.chat_history),
        "chat_history.pdf",
        "application/pdf"
    )
    st.download_button(
        "Download Chat History as DOCX",
        create_docx(st.session_state.chat_history),
        "chat_history.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Email and WhatsApp sharing
recipient_email = st.text_input("Recipient Email")
if st.button("Send Email"):
    content = "\n".join(
        [f"{'You' if chat['role'] == 'user' else 'Bot'}: {chat['message']}" for chat in st.session_state.chat_history]
    )
    send_email_via_smtp(content, recipient_email)

whatsapp_link = generate_whatsapp_link(st.session_state.chat_history)
st.markdown(f"[Share via WhatsApp]({whatsapp_link})", unsafe_allow_html=True)
